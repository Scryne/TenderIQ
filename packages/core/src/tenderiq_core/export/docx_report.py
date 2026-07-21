"""Word rapor üretimi (Sprint 3.2, §4.1) — python-docx, `tenderiq-report-template`.

Şablon: kapak bilgileri (firma + ihale + tarih, opsiyonel logo), kategori
başına tablo (içerik kolonları + İnceleme + Kaynak işareti) ve sonda numaralı
"Kaynaklar" bölümü. python-docx gerçek dipnot API'si sunmadığından kaynak
referansları endnote desenidir: tabloda ``[n]`` işareti → Kaynaklar'da
``[n] doküman, s. X · bölüm — "alıntı"`` (sayfa no her referansta görünür).

python-docx yalnız bu modülde ve lazy import edilir (`export` extra'sı
kurulmayan süreçler modülü import edebilir).
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from tenderiq_core.export.report import (
    REVIEW_STATUS_LABELS,
    SourceRef,
    TenderReport,
    truncate_quote,
)


def build_docx_report(report: TenderReport, *, logo: bytes | None = None) -> bytes:
    """Rapordan Word (.docx) baytları üretir (saf — DB/IO yok)."""
    from docx import Document  # opsiyonel bağımlılık (export extra) — lazy
    from docx.shared import Inches, Pt

    doc = Document()

    if logo is not None:
        doc.add_picture(BytesIO(logo), width=Inches(1.6))

    doc.add_heading("İhale Analiz Raporu", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Firma: {report.organization}\n").bold = True
    meta.add_run(f"İhale: {report.tender_title}\n")
    meta.add_run(f"Oluşturma: {report.generated_at.strftime('%d.%m.%Y %H:%M')}\n")
    meta.add_run("Bu rapor TenderIQ ile insan onaylı analizden üretilmiştir.").italic = True

    # Endnote deseni: tablo hücresindeki [n] işareti sondaki Kaynaklar'a bağlanır.
    references: list[SourceRef] = []

    for section in report.sections:
        if not section.items:
            continue
        doc.add_heading(f"{section.title} ({len(section.items)})", level=1)
        table = doc.add_table(rows=1, cols=len(section.headers) + 2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for column, header in enumerate((*section.headers, "İnceleme", "Kaynak")):
            run = header_cells[column].paragraphs[0].add_run(header)
            run.bold = True
        for item in section.items:
            references.append(item.source)
            row_cells = table.add_row().cells
            for column, cell_text in enumerate(item.cells):
                row_cells[column].text = cell_text
            row_cells[len(item.cells)].text = REVIEW_STATUS_LABELS[item.review_status]
            row_cells[len(item.cells) + 1].text = f"[{len(references)}]"

    if references:
        doc.add_heading("Kaynaklar", level=1)
        for number, source in enumerate(references, start=1):
            paragraph = doc.add_paragraph()
            marker = paragraph.add_run(f"[{number}] ")
            marker.bold = True
            paragraph.add_run(f"{source.location()} — ")
            quote = paragraph.add_run(f"“{truncate_quote(source.quote)}”")
            quote.italic = True
            _set_font_size(paragraph, Pt(9))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_font_size(paragraph: Any, size: Any) -> None:
    for run in paragraph.runs:
        run.font.size = size
