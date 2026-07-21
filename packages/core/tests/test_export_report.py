"""Export builder birim testleri (Sprint 3.2, §4.1): rapor modeli → docx/xlsx.

Builder'lar saftır (DB/IO yok); üretilen baytlar python-docx/openpyxl ile geri
açılıp içerik sözleşmesi doğrulanır: tablo kolonları, inceleme etiketi, [n]
kaynak işareti ve sayfa numaralı Kaynaklar bölümü (citation-first, ADR-0006).
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from tenderiq_core.export import (
    ReportItem,
    ReportSection,
    SourceRef,
    TenderReport,
    build_docx_report,
    build_xlsx_report,
    truncate_quote,
)
from tenderiq_core.findings import ReviewStatus


def _report() -> TenderReport:
    source = SourceRef(
        document="şartname.pdf",
        page=4,
        section="Madde 7.2",
        quote="Yüklenici tüm maddeleri karşılamak zorundadır.",
    )
    return TenderReport(
        organization="Örnek A.Ş.",
        tender_title="Donanım Alımı",
        generated_at=datetime(2026, 7, 21, 14, 30),
        sections=(
            ReportSection(
                title="Gereksinimler",
                headers=("Gereksinim", "Tür", "Zorunlu"),
                items=(
                    ReportItem(
                        cells=("Yüklenici tüm maddeleri karşılamalıdır.", "İdari", "Evet"),
                        review_status=ReviewStatus.APPROVED,
                        source=source,
                    ),
                    ReportItem(
                        cells=("Teminat mektubu sunulmalıdır.", "Mali", "Evet"),
                        review_status=ReviewStatus.EDITED,
                        source=source,
                    ),
                ),
            ),
            # Boş bölüm: docx'te hiç çizilmez, xlsx'te sayfa açılmaz.
            ReportSection(title="Riskler", headers=("Risk Maddesi", "Önem", "Kategori"), items=()),
        ),
    )


def test_truncate_quote_kisa_metni_degistirmez() -> None:
    assert truncate_quote("kısa alıntı") == "kısa alıntı"


def test_truncate_quote_uzun_metni_kelime_sinirinda_keser() -> None:
    quote = "kelime " * 100
    result = truncate_quote(quote, max_chars=50)
    assert len(result) <= 51  # 50 + '…'
    assert result.endswith("…")
    assert " kelime…" not in result or not result.endswith(" …")


def test_is_empty_yalniz_bos_bolumlerde_true() -> None:
    assert not _report().is_empty()
    empty = TenderReport(
        organization="X",
        tender_title="Y",
        generated_at=datetime(2026, 7, 21),
        sections=(ReportSection(title="Gereksinimler", headers=("A",), items=()),),
    )
    assert empty.is_empty()


def test_docx_tablolar_kaynakca_ve_inceleme_etiketi() -> None:
    from docx import Document

    data = build_docx_report(_report())
    document = Document(BytesIO(data))

    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert any("Gereksinimler (2)" in h for h in headings)
    assert "Kaynaklar" in headings
    assert all("Riskler" not in h for h in headings)  # boş bölüm çizilmedi

    assert len(document.tables) == 1
    table = document.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == [
        "Gereksinim",
        "Tür",
        "Zorunlu",
        "İnceleme",
        "Kaynak",
    ]
    first_row = [cell.text for cell in table.rows[1].cells]
    assert first_row[0] == "Yüklenici tüm maddeleri karşılamalıdır."
    assert first_row[3] == "Onaylandı"
    assert first_row[4] == "[1]"
    assert [cell.text for cell in table.rows[2].cells][3] == "Düzeltildi"

    text = "\n".join(p.text for p in document.paragraphs)
    assert "şartname.pdf, s. 4 · Madde 7.2" in text  # sayfa no'lu kaynak referansı
    assert "Örnek A.Ş." in text
    assert "Donanım Alımı" in text


def test_xlsx_ozet_ve_kategori_sayfasi_kaynak_kolonlari() -> None:
    from openpyxl import load_workbook

    data = build_xlsx_report(_report())
    workbook = load_workbook(BytesIO(data))

    assert workbook.sheetnames == ["Özet", "Gereksinimler"]  # boş Riskler sayfası yok
    summary = workbook["Özet"]
    assert summary["B3"].value == "Örnek A.Ş."
    assert summary["B4"].value == "Donanım Alımı"

    sheet = workbook["Gereksinimler"]
    header = [cell.value for cell in sheet[1]]
    assert header == [
        "Gereksinim",
        "Tür",
        "Zorunlu",
        "İnceleme",
        "Doküman",
        "Sayfa",
        "Bölüm",
        "Alıntı",
    ]
    row = [cell.value for cell in sheet[2]]
    assert row[0] == "Yüklenici tüm maddeleri karşılamalıdır."
    assert row[3] == "Onaylandı"
    assert row[4] == "şartname.pdf"
    assert row[5] == 4  # sayfa no her satırda
    assert row[6] == "Madde 7.2"
